import streamlit as st
from pypdf import PdfMerger, PdfReader, PdfWriter
from io import BytesIO
from docx import Document
import tempfile
import os
# st.sortable_items kullanabilmek için gerekli import
from streamlit_sortable import sortable_items 

# docx2pdf'i koşullu olarak import ediyoruz, yoksa hata vermez
try:
    # docx2pdf kütüphanesi harici bir uygulama (Word/LibreOffice) gerektirir!
    import docx2pdf 
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    
st.set_page_config(page_title="Belge Birleştirici", page_icon="📎", layout="centered")
st.title("📎 PDF & Word Birleştirici - Streamlit")
st.markdown("Bu uygulama PDF ve Word (DOCX) dosyalarını yükleyip sürükle-bırak yöntemiyle sırasını belirleyerek tek bir dosya haline getirir.")

st.markdown("---")

uploaded_files = st.file_uploader(
    "PDF veya Word dosyalarını yükleyin (çoklu seçim mümkün)",
    type=["pdf", "docx"],
    accept_multiple_files=True
)

# Dosya yüklenmemişse bilgilendirme ve çıkış
if not uploaded_files:
    st.info("Başlamak için PDF veya Word dosyalarını yükleyin.")
    st.markdown("---")
    st.caption("Not: Çok büyük dosyalarda bellek sınırları sorun oluşturabilir. Yerel çalıştırma daha stabil olabilir.")
    st.markdown("""
    **Gereksinimler**:
    - `pip install streamlit`
    - `pip install pypdf`
    - `pip install python-docx`
    - `pip install streamlit-sortable`
    - **DOCX+PDF birleştirme için**: `pip install docx2pdf` (Sisteminizde **Microsoft Word** veya **LibreOffice** kurulu olmalıdır, çünkü `docx2pdf` arkaplanda bu programları kullanır.)
    
    **Çalıştırma**:
    ```
    streamlit run combine.py
    ```
    """)
    st.stop() # Kodun geri kalanını çalıştırmayı durdur

# --- Dosya Sıralama ---
file_names = [f.name for f in uploaded_files]
st.subheader("Dosya sırası (sürükleyerek değiştirin)")

# Dosya adlarını sırala
sorted_file_names = sortable_items(file_names, key="file_sort")
# Orijinal dosya nesnelerini sıralanmış listeye dönüştür
sorted_files = [uploaded_files[file_names.index(name)] for name in sorted_file_names]

st.markdown("---")

# --- PDF Sayfa Silme / Taşıma ---
pdf_files_in_list = [n for n in file_names if n.lower().endswith('.pdf')]
if pdf_files_in_list:
    st.subheader("📄 PDF Sayfa Yönetimi")
    pdf_manage_name = st.selectbox("Sayfa yönetimi için bir PDF seçin", pdf_files_in_list)

    if pdf_manage_name:
        try:
            pdf_file = uploaded_files[file_names.index(pdf_manage_name)]
            # Dosyayı okumadan önce başlangıca git
            pdf_file.seek(0)
            reader = PdfReader(pdf_file)
            total_pages = len(reader.pages)

            st.write(f"Toplam sayfa: **{total_pages}**")

            page_list = [f"Sayfa {i+1}" for i in range(total_pages)]
            st.write("Sayfaları sürükleyerek yeniden sıralayın veya seçerek silin.")

            reordered = sortable_items(page_list, key=f"sort_pages_{pdf_manage_name}")
            delete_pages = st.multiselect("Silinecek sayfalar", reordered)

            if st.button("📌 Yeni PDF Üret (Sayfa Silme / Taşıma)"):
                writer = PdfWriter()
                for page_name in reordered:
                    idx = int(page_name.split()[1]) - 1
                    # Silinecekler listesinde yoksa sayfayı ekle
                    if page_name not in delete_pages:
                        writer.add_page(reader.pages[idx])

                out_pdf = BytesIO()
                writer.write(out_pdf)
                out_pdf.seek(0)

                st.success("Yeni PDF oluşturuldu!")
                st.download_button(
                    "📥 Düzenlenmiş PDF'i İndir",
                    out_pdf,
                    f"edited_{pdf_manage_name}",
                    mime="application/pdf",
                )
        except Exception as e:
            st.error(f"PDF Sayfa Yönetimi Hatası: {e}")
            
st.markdown("---")

# --- PDF Birleştirme ---
pdf_files_to_merge = [file for file in sorted_files if file.name.lower().endswith(".pdf")]
if st.button("🔀 PDF'leri Birleştir", disabled=not pdf_files_to_merge):
    if not pdf_files_to_merge:
        st.error("Birleştirilecek PDF dosyası bulunamadı.")
    else:
        try:
            merger = PdfMerger()
            for file in pdf_files_to_merge:
                file.seek(0) # Dosyayı tekrar en başa al
                merger.append(file)
            out = BytesIO()
            merger.write(out)
            merger.close()
            out.seek(0)

            st.success("PDF başarıyla birleştirildi!")
            st.download_button(
                label="📥 Birleşmiş PDF'i İndir",
                data=out,
                file_name="merged.pdf",
                mime="application/pdf",
            )
        except Exception as e:
            st.error(f"PDF birleştirme hatası: {e}")

# --- Word Birleştirme (DOCX -> DOCX) ---
word_files_to_merge = [file for file in sorted_files if file.name.lower().endswith(".docx")]
if st.button("📝 Word (DOCX) Birleştir", disabled=not word_files_to_merge):
    if not word_files_to_merge:
        st.error("Birleştirilecek Word dosyası bulunamadı.")
    else:
        try:
            merged_doc = Document()
            first = True
            
            # Geçici dosyaları izlemek için liste
            temp_files_to_clean = [] 
            
            for file in word_files_to_merge:
                # Geçici dosyayı oluştur
                temp_path = tempfile.mktemp(suffix=".docx")
                temp_files_to_clean.append(temp_path)
                
                # Yüklenen dosyayı geçici dosyaya yaz
                file.seek(0)
                with open(temp_path, "wb") as tmp:
                    tmp.write(file.getbuffer())
                
                # Geçici dosyayı oku
                sub_doc = Document(temp_path)

                # İlk belge değilse sayfa sonu ekle
                if not first:
                    merged_doc.add_page_break()
                
                # Paragrafları birleştir (stil koruması dene)
                for p in sub_doc.paragraphs:
                    # Not: Bu yöntem, tablolar veya resimler gibi karmaşık Word yapılarını korumaz.
                    merged_doc.add_paragraph(p.text, style=p.style) 
                
                # İlk belge eklendi
                first = False

            # Birleşmiş belgeyi kaydet ve indir
            out_docx = BytesIO()
            merged_doc.save(out_docx)
            out_docx.seek(0)

            st.success("Word belgeleri birleştirildi!")
            st.download_button(
                "📥 Birleşmiş Word Belgesini İndir",
                out_docx,
                "merged.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            
            # Temizleme
            for path in temp_files_to_clean:
                if os.path.exists(path):
                    os.remove(path)
                    
        except Exception as e:
            st.error(f"Word birleştirme hatası: {e}")
            
# --- DOCX + PDF -> Tek PDF Birleştirme ---
if DOCX2PDF_AVAILABLE:
    if st.button("📄 DOCX + PDF → Tek PDF Birleştir", disabled=(not pdf_files_to_merge and not word_files_to_merge)):
        if not pdf_files_to_merge and not word_files_to_merge:
            st.error("PDF veya DOCX bulunamadı.")
        else:
            try:
                temp_pdf_list = []
                temp_files_to_clean = [] # Tüm geçici dosyalar için

                # DOCX → PDF dönüşümü
                docx_files_to_convert = [f for f in sorted_files if f.name.lower().endswith(".docx")]
                st.info(f"Dönüştürülüyor: {len(docx_files_to_convert)} DOCX dosyası PDF'e çevriliyor...")
                
                for file in docx_files_to_convert:
                    tmp_docx = tempfile.mktemp(suffix=".docx")
                    tmp_pdf = tempfile.mktemp(suffix=".pdf")
                    temp_files_to_clean.extend([tmp_docx, tmp_pdf])
                    
                    file.seek(0)
                    with open(tmp_docx, "wb") as tmp:
                        tmp.write(file.getbuffer())
                        
                    # docx2pdf harici bir uygulama gerektirir (Word veya LibreOffice)
                    docx2pdf.convert(tmp_docx, tmp_pdf)
                    temp_pdf_list.append(tmp_pdf)

                # PDF birleştirme
                merger = PdfMerger()
                pdf_index = 0
                for file in sorted_files:
                    if file.name.lower().endswith(".pdf"):
                        file.seek(0)
                        merger.append(file)
                    else: # DOCX'ten dönüştürülmüş PDF'i ekle (sırayı koruyarak)
                        merger.append(temp_pdf_list[pdf_index])
                        pdf_index += 1 

                out = BytesIO()
                merger.write(out)
                merger.close()
                out.seek(0)

                st.success("DOCX + PDF birlikte tek PDF olarak birleştirildi!")
                st.download_button(
                    "📥 Tek PDF Olarak İndir",
                    out,
                    "merged_all.pdf",
                    mime="application/pdf",
                )
                
                # Temizleme
                for path in temp_files_to_clean:
                    if os.path.exists(path):
                        os.remove(path)
                        
            except Exception as e:
                st.error(f"Birleştirme hatası: {e}")
                st.error("DOCX'ten PDF'e dönüştürme için sisteminizde **Microsoft Word** veya **LibreOffice** kurulu olmalıdır.")
else:
    st.warning("⚠️ `docx2pdf` modülü bulunamadı. DOCX + PDF birleştirme işlevi devre dışı.")


st.markdown("---")
st.caption("Not: Çok büyük dosyalarda bellek sınırları sorun oluşturabilir. Yerel çalıştırma daha stabil olabilir.")
st.markdown("""
**Gereksinimler**:
- `pip install streamlit`
- `pip install pypdf`
- `pip install python-docx`
- `pip install streamlit-sortable`
- **DOCX+PDF birleştirme için**: `pip install docx2pdf` (Sisteminizde **Microsoft Word** veya **LibreOffice** kurulu olmalıdır, çünkü `docx2pdf` arkaplanda bu programları kullanır.)

**Çalıştırma**:
""") 👈 **Bu satırda sadece üç tırnak işareti (""") olmalıdır.**
