import streamlit as st
import pdfplumber
import pandas as pd
from io import BytesIO
from docx import Document

def make_unique_columns(columns):
    seen = {}
    result = []
    for col in columns:
        if col not in seen:
            seen[col] = 0
            result.append(col)
        else:
            seen[col] += 1
            result.append(f"{col}_{seen[col]}")
    return result

st.title("PDF Dönüştürücü")

uploaded_file = st.file_uploader("PDF dosyanızı yükleyin", type=["pdf"])

if uploaded_file:
    option = st.selectbox("Çevirme seçeneği:", ["Excel", "Word"])

    if option == "Excel":
        all_tables = []
        with pdfplumber.open(uploaded_file) as pdf:
            for i, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                if tables:
                    for j, table in enumerate(tables):
                        df = pd.DataFrame(table[1:], columns=table[0])
                        df.columns = make_unique_columns(df.columns)
                        all_tables.append((f"Sayfa {i+1} Tablo {j+1}", df))

        if all_tables:
            export_option = st.radio("Excel dosyasını nasıl oluşturmak istersiniz?", 
                                     ("Her tablo ayrı sayfada", "Tüm tablolar tek sayfada"))

            output = BytesIO()
            with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                if export_option == "Her tablo ayrı sayfada":
                    for sheet_name, df in all_tables:
                        safe_sheet_name = sheet_name[:31]
                        df.to_excel(writer, sheet_name=safe_sheet_name, index=False)
                else:
                    worksheet_name = "Tüm Tablolar"
                    worksheet = writer.book.add_worksheet(worksheet_name)
                    writer.sheets[worksheet_name] = worksheet
                    startrow = 0
                    for title, df in all_tables:
                        worksheet.write(startrow, 0, title)
                        df.to_excel(writer, sheet_name=worksheet_name, startrow=startrow + 1, index=False, header=True)
                        startrow += len(df) + 3

            processed_data = output.getvalue()

            st.download_button(
                label="Excel Dosyasını İndir",
                data=processed_data,
                file_name="pdf_tablosu.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        else:
            st.warning("PDF'de tablo bulunamadı veya PDF taranmış bir görüntü içeriyor.")

    elif option == "Word":
        doc = Document()
        with pdfplumber.open(uploaded_file) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    doc.add_paragraph(f"Sayfa {i+1}")
                    doc.add_paragraph(text)
                tables = page.extract_tables()
                for table in tables:
                    t = doc.add_table(rows=1, cols=len(table[0]))
                    hdr_cells = t.rows[0].cells
                    for j, heading in enumerate(table[0]):
                        hdr_cells[j].text = heading
                    for row in table[1:]:
                        row_cells = t.add_row().cells
                        for k, cell in enumerate(row):
                            row_cells[k].text = str(cell)
                    doc.add_paragraph()

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        st.download_button(
            label="Word Dosyasını İndir",
            data=output,
            file_name="converted.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
